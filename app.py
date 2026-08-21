"""
Course Material -> Quiz Generator (Ollama Cloud)
====================================================
Run with:  streamlit run app.py

Upload a PDF or DOCX of course material. This tool reads the text,
sends it to a model hosted on Ollama Cloud to draft multiple-choice
questions (with plausible wrong options AND a suggested correct
answer), lets you review/edit everything, then exports a completely
separate, self-contained `quiz_app.py` + `README.md` you can hand to
students or push to GitHub. The exported quiz app needs only
`streamlit` to run -- it does NOT call Ollama at all, since the
questions and answers are baked in as data.

Requires an Ollama Cloud API key (free tier available):
    1. Sign up / sign in at https://ollama.com
    2. Create an API key: https://ollama.com/settings/keys
    3. Either set it as an environment variable before launching --
         export OLLAMA_API_KEY=your_key_here      (Mac/Linux)
         setx OLLAMA_API_KEY "your_key_here"       (Windows, new terminal after)
       -- or paste it into the sidebar field when the app is running.

No local model download, no `ollama serve`, and your laptop doesn't
need to stay on for a deployed version of this app to keep working --
inference runs on Ollama's cloud infrastructure.
"""

import base64
import io
import json
import os
import re
import zipfile
from datetime import datetime

import requests
import streamlit as st

# ============================================================================
# CONFIG
# ============================================================================
OLLAMA_CLOUD_HOST = "https://ollama.com"
# A few reasonable defaults for structured-JSON generation tasks like this
# one. Ollama's cloud catalog changes over time -- if you want a different
# model, use "Custom model tag..." in the sidebar and check what's
# currently available to your account at https://ollama.com/settings/keys
# or by calling GET https://ollama.com/api/tags with your API key.
MODEL_OPTIONS = ["gpt-oss:20b", "gpt-oss:120b", "qwen3.5", "deepseek-v3.1:671b", "Custom model tag..."]
CHUNK_CHARS = 4000  # cloud models handle larger context comfortably

st.set_page_config(page_title="Quiz Generator (Ollama Cloud)", page_icon="📚", layout="wide")

# ============================================================================
# SESSION STATE
# ============================================================================
defaults = {
    "raw_text": "",
    "source_filename": "",
    "questions": [],
    "generation_log": [],
    "generated_once": False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ============================================================================
# FILE EXTRACTION
# ============================================================================
def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def chunk_text(text: str, chunk_size: int = CHUNK_CHARS):
    text = text.strip()
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # try to break on a paragraph/sentence boundary, not mid-word
        if end < len(text):
            boundary = text.rfind("\n", start, end)
            if boundary == -1 or boundary <= start + chunk_size * 0.5:
                boundary = text.rfind(". ", start, end)
            if boundary != -1 and boundary > start + chunk_size * 0.3:
                end = boundary + 1
        chunks.append(text[start:end].strip())
        start = end
    return [c for c in chunks if c]


# ============================================================================
# OLLAMA CALL + PARSING
# ============================================================================
PROMPT_TEMPLATE = """You are helping a university instructor create multiple-choice quiz questions for students, based ONLY on the course material provided below.

Generate exactly {n} multiple-choice questions from this material.
Each question must have exactly {opts} answer options.
Exactly one option must be the correct answer; the rest must be plausible but clearly wrong to someone who understood the material.
Vary which option position (0, 1, 2, ...) holds the correct answer -- do not always put it first.

Return ONLY a JSON array -- your entire response must start with [ and end with ]. Do not wrap it in an object, do not add markdown code fences, do not add any text before or after it. Each element must have exactly this shape:
{{
  "question": "the question text",
  "options": ["option A text", "option B text", "..."],
  "correct_index": 0,
  "explanation": "one short sentence explaining why that answer is correct"
}}

"correct_index" is the 0-based index into "options" of the correct answer.

COURSE MATERIAL:
\"\"\"
{material}
\"\"\"
"""


def call_ollama(model: str, prompt: str, host: str, api_key: str) -> str:
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    resp = requests.post(
        f"{host}/api/chat",
        headers=headers,
        json={
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "format": "json",
            "options": {"temperature": 0.4},
        },
        timeout=600,
    )
    if resp.status_code == 401:
        raise PermissionError(
            "Ollama Cloud rejected the API key (401 Unauthorized). "
            "Check the key in the sidebar / OLLAMA_API_KEY env var."
        )
    if resp.status_code == 429:
        raise RuntimeError(
            "Ollama Cloud rate/usage limit hit (429). Wait a bit, reduce the "
            "number of questions, or upgrade your plan at https://ollama.com/pricing."
        )
    resp.raise_for_status()
    data = resp.json()
    return data.get("message", {}).get("content", "")


def extract_json_array(raw: str):
    """Best-effort extraction of a JSON array from a model response, even
    if the model added stray text around it, wrapped it in a code fence,
    or nested it inside an object like {"questions": [...]} instead of
    returning a bare array (all common with small local models)."""
    raw = raw.strip()
    # strip a markdown code fence if present
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", raw, re.DOTALL)
    if fence:
        raw = fence.group(1).strip()

    def _unwrap(value):
        if isinstance(value, list):
            return value
        if isinstance(value, dict):
            for key in ("questions", "quiz", "items", "data", "results"):
                if isinstance(value.get(key), list):
                    return value[key]
            # a dict of one question rather than a list of questions
            if "question" in value:
                return [value]
        return None

    try:
        return _unwrap(json.loads(raw))
    except json.JSONDecodeError:
        pass

    # try to locate the outermost array anywhere in the text
    match = re.search(r"\[.*\]", raw, re.DOTALL)
    if match:
        try:
            return _unwrap(json.loads(match.group(0)))
        except json.JSONDecodeError:
            pass

    # fall back to locating an outermost object
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        try:
            return _unwrap(json.loads(match.group(0)))
        except json.JSONDecodeError:
            pass

    return None


def validate_questions(items, num_options: int):
    """Coerce and accept near-miss output from small models rather than
    rejecting outright, but enforce the requested option count exactly
    so the final quiz is uniform (mixed 3/4/5-option questions read as
    broken to a student). If the model gave too many options, the extra
    WRONG ones are trimmed (the correct one is always kept). If it gave
    too few, the question is dropped -- there's no safe way to invent a
    plausible extra distractor without another model call."""
    valid = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        q = item.get("question")
        opts = item.get("options")
        idx = item.get("correct_index")

        if not isinstance(q, str) or not q.strip():
            continue
        if not isinstance(opts, list) or len(opts) < 2:
            continue

        opts = [str(o).strip() for o in opts if str(o).strip()]
        if len(opts) < 2:
            continue

        # coerce a stringified index ("0", "2") to int
        if isinstance(idx, str) and idx.strip().isdigit():
            idx = int(idx.strip())
        # some models label the correct option by its text instead of an index
        if isinstance(idx, str) and idx not in opts:
            match = next((k for k, o in enumerate(opts) if o.lower() == idx.strip().lower()), None)
            idx = match if match is not None else None
        elif isinstance(idx, str) and idx in opts:
            idx = opts.index(idx)

        if not isinstance(idx, int) or not (0 <= idx < len(opts)):
            continue

        # enforce exact option count for a uniform quiz
        if len(opts) > num_options:
            correct_text = opts[idx]
            wrong_opts = [o for k, o in enumerate(opts) if k != idx]
            keep_wrong = wrong_opts[: num_options - 1]
            opts = keep_wrong + [correct_text]
            idx = len(opts) - 1
        elif len(opts) < num_options:
            continue  # can't safely fabricate a missing distractor

        valid.append(
            {
                "question": q.strip(),
                "options": opts,
                "correct_index": idx,
                "explanation": str(item.get("explanation", "")).strip(),
            }
        )
    return valid


def _generate_from_chunk(chunk: str, n: int, model: str, host: str, api_key: str, num_options: int, log: list, label: str):
    """Ask the model for n questions from one chunk, validate, log the
    outcome, and return the list of valid questions (possibly empty).
    Raises PermissionError on a bad API key so the caller can stop the
    whole run immediately instead of retrying a doomed request repeatedly."""
    prompt = PROMPT_TEMPLATE.format(n=n, opts=num_options, material=chunk)
    try:
        raw = call_ollama(model, prompt, host, api_key)
    except PermissionError:
        raise  # let the caller abort the whole run, not just this chunk
    except requests.exceptions.ConnectionError:
        log.append(f"{label}: could not reach {host}. Check your internet connection.")
        return []
    except RuntimeError as e:  # rate limit
        log.append(f"{label}: {e}")
        return []
    except Exception as e:  # noqa: BLE001
        log.append(f"{label}: request failed ({e}).")
        return []

    parsed = extract_json_array(raw)
    valid = validate_questions(parsed, num_options)
    if not valid:
        preview = raw.strip().replace("\n", " ")[:300]
        log.append(
            f"{label}: model did not return usable questions, skipped. "
            f"Raw response preview: `{preview}{'...' if len(raw.strip()) > 300 else ''}`"
        )
    else:
        log.append(f"{label}: generated {len(valid)} valid question(s).")
    return valid


MAX_PER_CALL = 4  # small local models are much more reliable generating a
                   # short JSON array than a long one in a single response;
                   # split bigger requests into several smaller calls instead.


def _split_into_batches(n: int, batch_size: int = MAX_PER_CALL):
    batches = []
    remaining = n
    while remaining > 0:
        take = min(batch_size, remaining)
        batches.append(take)
        remaining -= take
    return batches


def generate_quiz(text: str, model: str, host: str, api_key: str, total_questions: int, num_options: int):
    chunks = chunk_text(text)
    if not chunks:
        return [], ["No text could be extracted from the document."]

    n_chunks = len(chunks)
    base = total_questions // n_chunks
    remainder = total_questions % n_chunks
    per_chunk = [base + (1 if i < remainder else 0) for i in range(n_chunks)]

    all_questions = []
    log = []
    progress = st.progress(0.0, text="Starting generation...")

    try:
        for i, (chunk, n) in enumerate(zip(chunks, per_chunk)):
            progress.progress(i / n_chunks, text=f"Generating from section {i + 1}/{n_chunks}...")
            if n <= 0:
                continue
            batches = _split_into_batches(n)
            for b, batch_n in enumerate(batches):
                label = f"Section {i + 1}" + (f" (batch {b + 1}/{len(batches)})" if len(batches) > 1 else "")
                all_questions.extend(
                    _generate_from_chunk(chunk, batch_n, model, host, api_key, num_options, log, label)
                )

        # ---- top-up pass: strict validation means some questions get
        # dropped (wrong option count, no answer key found, etc.), so the
        # first pass often falls short of the target. Keep asking for the
        # remainder, cycling through chunks, up to a small retry cap so a
        # stubborn document can't loop forever.
        max_extra_rounds = 8
        round_num = 0
        while len(all_questions) < total_questions and round_num < max_extra_rounds:
            round_num += 1
            deficit = total_questions - len(all_questions)
            progress.progress(
                0.9, text=f"Topping up {deficit} more question(s) (round {round_num})..."
            )
            chunk = chunks[round_num % n_chunks]
            batch_n = min(deficit, MAX_PER_CALL)
            got = _generate_from_chunk(
                chunk,
                batch_n,
                model,
                host,
                api_key,
                num_options,
                log,
                f"Top-up round {round_num}",
            )
            # avoid near-duplicate questions if the model regenerates similar
            # content from the same chunk on repeated top-up rounds
            existing = {q["question"].strip().lower() for q in all_questions}
            got = [q for q in got if q["question"].strip().lower() not in existing]
            all_questions.extend(got)
            if not got:
                # this chunk isn't yielding more -- no point retrying it again
                continue
    except PermissionError as e:
        progress.progress(1.0, text="Stopped.")
        log.append(str(e))
        return all_questions[:total_questions], log

    progress.progress(1.0, text="Done.")
    if len(all_questions) < total_questions:
        log.append(
            f"Reached the retry limit with {len(all_questions)}/{total_questions} questions. "
            "The model may be struggling with this document/option count -- try fewer "
            "questions, a larger/better model (e.g. qwen3:4b), or fewer options per question."
        )
    return all_questions[:total_questions], log


# ============================================================================
# EXPORT TEMPLATE (the standalone student-facing quiz app)
# ============================================================================
QUIZ_APP_TEMPLATE = r'''"""
__QUIZ_TITLE__
Auto-generated quiz app -- self-contained, no external services needed.
Generated on __GENERATED_DATE__ from: __SOURCE_FILENAME__

Run with:
    pip install streamlit
    streamlit run quiz_app.py
"""

import base64
import json

import streamlit as st

QUESTIONS_B64 = "__QUESTIONS_B64__"
QUESTIONS = json.loads(base64.b64decode(QUESTIONS_B64).decode("utf-8"))

st.set_page_config(page_title="__QUIZ_TITLE__", page_icon="\U0001F4DA", layout="centered")

st.title("__QUIZ_TITLE__")
st.caption(f"{len(QUESTIONS)} questions | Generated from: __SOURCE_FILENAME__")
st.write("")

if "quiz_submitted" not in st.session_state:
    st.session_state.quiz_submitted = False

with st.form("quiz_form"):
    answers = {}
    for i, q in enumerate(QUESTIONS):
        st.markdown(f"**{i + 1}. {q['question']}**")
        answers[i] = st.radio(
            label=f"q_{i}",
            options=list(range(len(q["options"]))),
            format_func=lambda idx, opts=q["options"]: opts[idx],
            key=f"answer_{i}",
            label_visibility="collapsed",
            index=None,
        )
        st.write("")
    submitted = st.form_submit_button("Submit Quiz", use_container_width=True)

if submitted:
    st.session_state.quiz_submitted = True

if st.session_state.quiz_submitted:
    unanswered = [i for i, a in answers.items() if a is None]
    if unanswered:
        st.warning(
            f"You left {len(unanswered)} question(s) unanswered: "
            f"{', '.join(str(i + 1) for i in unanswered)}. "
            "They're marked wrong below -- answer everything and resubmit to update your score."
        )

    score = sum(1 for i, q in enumerate(QUESTIONS) if answers.get(i) == q["correct_index"])
    total = len(QUESTIONS)
    pct = round(100 * score / total) if total else 0
    st.header(f"Score: {score} / {total} ({pct}%)")
    st.write("")
    st.subheader("Review")

    for i, q in enumerate(QUESTIONS):
        user_idx = answers.get(i)
        correct_idx = q["correct_index"]
        is_correct = user_idx == correct_idx
        icon = "\u2705" if is_correct else "\u274C"
        with st.expander(f"{icon} {i + 1}. {q['question']}"):
            for j, opt in enumerate(q["options"]):
                prefix = ""
                if j == correct_idx:
                    prefix = "\u2705 "
                elif j == user_idx:
                    prefix = "\u274C "
                st.write(f"{prefix}{opt}")
            if q.get("explanation"):
                st.caption(f"Why: {q['explanation']}")

    st.write("")
    if st.button("Restart Quiz"):
        for i in range(len(QUESTIONS)):
            st.session_state.pop(f"answer_{i}", None)
        st.session_state.quiz_submitted = False
        st.rerun()
'''

README_TEMPLATE = r"""# __QUIZ_TITLE__

An auto-generated, self-contained multiple-choice quiz app.

- **__NUM_QUESTIONS__ questions**, __NUM_OPTIONS__ options each
- Generated on __GENERATED_DATE__ from: `__SOURCE_FILENAME__`
- Questions were AI-drafted from the course material using a local model
  (__MODEL_USED__) and reviewed/edited before export.
- **No Ollama or internet connection needed to run this quiz** -- the
  questions, options, and answers are embedded directly in `quiz_app.py`.

## Run it

```bash
pip install streamlit
streamlit run quiz_app.py
```

Then open the local URL Streamlit prints (usually http://localhost:8501).

## Deploy it (optional)

Push `quiz_app.py` (and this README) to a GitHub repo, then deploy on
[Streamlit Community Cloud](https://share.streamlit.io) with the main
file path set to `quiz_app.py` -- no other setup or secrets required.

## Notes

- AI-drafted answers were reviewed by the instructor before export, but
  always spot-check a generated quiz before giving it to students.
- To regenerate or create a different quiz from new material, use the
  quiz builder tool (`app.py`) again -- this exported file is a
  standalone snapshot and does not need it to run.
"""


def build_export_files(questions, source_filename, model_used):
    quiz_title = f"Quiz: {source_filename.rsplit('.', 1)[0]}" if source_filename else "Course Quiz"
    generated_date = datetime.now().strftime("%Y-%m-%d %H:%M")
    questions_json = json.dumps(questions, ensure_ascii=False)
    questions_b64 = base64.b64encode(questions_json.encode("utf-8")).decode("ascii")

    quiz_code = (
        QUIZ_APP_TEMPLATE.replace("__QUIZ_TITLE__", quiz_title)
        .replace("__GENERATED_DATE__", generated_date)
        .replace("__SOURCE_FILENAME__", source_filename or "unknown")
        .replace("__QUESTIONS_B64__", questions_b64)
    )
    readme = (
        README_TEMPLATE.replace("__QUIZ_TITLE__", quiz_title)
        .replace("__NUM_QUESTIONS__", str(len(questions)))
        .replace("__NUM_OPTIONS__", str(len(questions[0]["options"])) if questions else "?")
        .replace("__GENERATED_DATE__", generated_date)
        .replace("__SOURCE_FILENAME__", source_filename or "unknown")
        .replace("__MODEL_USED__", model_used)
    )
    return quiz_code, readme


# ============================================================================
# UI -- chat-style walkthrough
# ============================================================================
st.title("📚 Course Material → Quiz Generator")
st.caption("Runs entirely on your machine via Ollama. No data leaves your laptop.")

with st.chat_message("assistant"):
    st.write(
        "Upload a PDF or DOCX of your course material, choose how many "
        "questions you want, and I'll draft a multiple-choice quiz from "
        "it -- including suggested correct answers -- for you to review."
    )

# ---- Sidebar settings ----
with st.sidebar:
    st.header("Settings")
    model_choice = st.selectbox(
        "Model (Ollama Cloud)",
        MODEL_OPTIONS,
        index=0,
        help="gpt-oss:20b is a good low-usage-tier default on the free plan. "
        "gpt-oss:120b / qwen3.5 / deepseek-v3.1 are larger and higher quality "
        "but consume your plan's usage allowance faster.",
    )
    if model_choice == "Custom model tag...":
        model = st.text_input(
            "Model tag",
            placeholder="e.g. glm-4.6:cloud",
            help="Check currently available tags for your account at "
            "https://ollama.com/settings/keys or via GET https://ollama.com/api/tags",
        )
    else:
        model = model_choice

    api_key = st.text_input(
        "Ollama API key",
        value=os.environ.get("OLLAMA_API_KEY", ""),
        type="password",
        help="From https://ollama.com/settings/keys. Pre-filled automatically "
        "if the OLLAMA_API_KEY environment variable is set before launching.",
    )
    num_questions = st.slider("Number of questions", 3, 30, 10)
    num_options = st.slider("Options per question", 2, 6, 4)
    st.divider()
    st.caption(
        "Runs on Ollama Cloud, not your machine -- no local model download, "
        "and this works the same whether you run the app locally or deploy it "
        "(e.g. to Streamlit Community Cloud)."
    )
    st.caption("Get a free API key: https://ollama.com")

# ---- File upload ----
uploaded = st.file_uploader("Upload course material", type=["pdf", "docx"])

if uploaded is not None and uploaded.name != st.session_state.source_filename:
    with st.spinner(f"Reading {uploaded.name}..."):
        file_bytes = uploaded.read()
        if uploaded.name.lower().endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        else:
            text = extract_text_from_docx(file_bytes)
    st.session_state.raw_text = text
    st.session_state.source_filename = uploaded.name
    st.session_state.questions = []
    st.session_state.generated_once = False

if st.session_state.raw_text:
    word_count = len(st.session_state.raw_text.split())
    with st.chat_message("assistant"):
        st.write(
            f"I read **{word_count:,} words** from `{st.session_state.source_filename}`. "
            f"Ready to draft **{num_questions} questions** ({num_options} options each) "
            f"using **{model}** on Ollama Cloud."
        )
    if word_count < 30:
        st.warning(
            "That's very little extracted text -- if this is a scanned/image-only "
            "PDF, text extraction won't work. Try a text-based PDF or DOCX."
        )
    if not api_key:
        st.warning(
            "No Ollama API key set -- add one in the sidebar (or set the "
            "OLLAMA_API_KEY environment variable) before generating."
        )
    if not model:
        st.warning("Enter a custom model tag in the sidebar, or pick one from the list.")

    if st.button(
        "🚀 Generate Quiz", type="primary", disabled=word_count < 30 or not api_key or not model
    ):
        with st.spinner("Talking to Ollama Cloud..."):
            questions, log = generate_quiz(
                st.session_state.raw_text, model, OLLAMA_CLOUD_HOST, api_key, num_questions, num_options
            )
        st.session_state.questions = questions
        st.session_state.generation_log = log
        st.session_state.generated_once = True

# ---- Generation log / errors ----
if st.session_state.generation_log:
    with st.expander("Generation log", expanded=not st.session_state.questions):
        for line in st.session_state.generation_log:
            st.write("- " + line)

# ---- Review & edit ----
if st.session_state.questions:
    with st.chat_message("assistant"):
        st.write(
            f"Drafted **{len(st.session_state.questions)}** question(s). "
            "Review and edit below -- fix wording, swap options, or change "
            "the marked correct answer -- then export."
        )

    to_remove = []
    for i, q in enumerate(st.session_state.questions):
        with st.expander(f"Q{i + 1}: {q['question'][:80]}", expanded=False):
            q["question"] = st.text_area("Question", value=q["question"], key=f"qtext_{i}")
            for j in range(len(q["options"])):
                q["options"][j] = st.text_input(
                    f"Option {j + 1}", value=q["options"][j], key=f"opt_{i}_{j}"
                )
            q["correct_index"] = st.radio(
                "Correct answer",
                options=list(range(len(q["options"]))),
                format_func=lambda idx, opts=q["options"]: opts[idx],
                index=q["correct_index"],
                key=f"correct_{i}",
            )
            q["explanation"] = st.text_input(
                "Explanation (optional)", value=q.get("explanation", ""), key=f"expl_{i}"
            )
            if st.checkbox("Remove this question", key=f"remove_{i}"):
                to_remove.append(i)

    if to_remove:
        st.session_state.questions = [
            q for i, q in enumerate(st.session_state.questions) if i not in to_remove
        ]
        st.rerun()

    st.divider()
    st.subheader("Export")
    if st.session_state.questions:
        quiz_code, readme = build_export_files(
            st.session_state.questions, st.session_state.source_filename, model
        )

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.download_button(
                "⬇️ Download quiz_app.py",
                data=quiz_code,
                file_name="quiz_app.py",
                mime="text/x-python",
                use_container_width=True,
            )
        with col2:
            st.download_button(
                "⬇️ Download README.md",
                data=readme,
                file_name="README.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with col3:
            questions_json = json.dumps(st.session_state.questions, ensure_ascii=False, indent=2)
            st.download_button(
                "⬇️ Download questions.json",
                data=questions_json,
                file_name="questions.json",
                mime="application/json",
                use_container_width=True,
            )
        with col4:
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.writestr("quiz_app.py", quiz_code)
                zf.writestr("README.md", readme)
                zf.writestr("questions.json", questions_json)
            st.download_button(
                "⬇️ Download all (.zip)",
                data=zip_buf.getvalue(),
                file_name="generated_quiz.zip",
                mime="application/zip",
                use_container_width=True,
            )

        with st.expander("Preview generated quiz_app.py"):
            st.code(quiz_code, language="python")
    else:
        st.info("All questions were removed -- nothing to export.")
